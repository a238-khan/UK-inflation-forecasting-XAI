from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd


def _project_root() -> Path:
    here = Path(__file__).resolve()
    return here.parents[1]


def _safe_read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(f"Missing file: {path}")
    return pd.read_csv(path)


def _fmt_heading(doc, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _set_uwe_style(doc):
    # Conservative "UWE-style" academic formatting:
    # - 2.54cm margins
    # - Times New Roman 12
    # - 1.5 line spacing
    # If you have an official UWE template, we can switch to it.
    from docx.shared import Cm, Pt

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Times New Roman"


def _set_paragraph_spacing(doc):
    from docx.shared import Pt

    for p in doc.paragraphs:
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_after = Pt(6)


def _add_df_table(doc, df: pd.DataFrame, title: str | None = None):
    if title:
        doc.add_paragraph(title)

    df = df.copy()
    # Ensure index is a column if meaningful
    if df.index.name or df.index.names != [None]:
        df = df.reset_index()

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, (float, np.floating)):
                cells[j].text = f"{val:.3f}"
            else:
                cells[j].text = str(val)


def _add_figure(doc, img_path: Path, caption: str, width_inches: float = 6.3):
    from docx.shared import Inches

    if not img_path.exists():
        return

    doc.add_picture(str(img_path), width=Inches(width_inches))
    doc.add_paragraph(caption)


def _sorted_figs(fig_paths: list[Path]) -> list[Path]:
    def key(p: Path):
        # sort by leading fig number if present
        name = p.stem
        digits = "".join([c for c in name if c.isdigit()])
        return (int(digits) if digits else 10**9, name)

    return sorted(fig_paths, key=key)


def _make_arima_plot(root: Path, out_dir: Path) -> Path:
    import matplotlib.pyplot as plt

    base = root / "results" / "baselines"
    dates = np.load(base / "arima_test_dates.npy", allow_pickle=True)
    actual = np.load(base / "arima_test_actual.npy", allow_pickle=True)
    forecast = np.load(base / "arima_test_forecast.npy", allow_pickle=True)

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "arima_actual_vs_forecast_test.png"

    plt.figure(figsize=(10, 4))
    plt.plot(dates, actual, label="Actual", color="black", linewidth=2.6)
    plt.plot(dates, forecast, label="ARIMA Forecast", color="tab:purple", linestyle="--", linewidth=2.2)
    plt.title("ARIMA Baseline — Actual vs Forecast (Test)")
    plt.xlabel("Date")
    plt.ylabel("CPI YoY")
    plt.grid(True, alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=220)
    plt.close()

    return out_path


def _make_rf_vs_hybrid_plot(root: Path, out_dir: Path) -> Path:
    import matplotlib.pyplot as plt
    import matplotlib.dates as mdates

    rf_dir = root / "results" / "baselines" / "rf"
    rf_pred_path = rf_dir / "rf_test_predictions.csv"
    rf_actual_path = rf_dir / "rf_test_actual.npy"
    rf_target_cols_path = rf_dir / "target_cols.json"

    rf_preds = pd.read_csv(rf_pred_path)
    rf_preds["date"] = pd.to_datetime(rf_preds["date"], errors="raise")
    rf_preds = rf_preds.sort_values("date").reset_index(drop=True)

    import json

    with open(rf_target_cols_path, "r", encoding="utf-8") as f:
        rf_target_cols = json.load(f)

    y_true = np.load(rf_actual_path)
    if y_true.ndim == 1:
        y_true = y_true.reshape(-1, len(rf_target_cols))

    rf_actuals = pd.DataFrame(y_true, columns=rf_target_cols)
    rf_actuals.insert(0, "date", rf_preds["date"].values)

    hybrid_dir = root / "results" / "hybrid"
    hybrid_files = {
        "t1": hybrid_dir / "xgb_cpi_t1_predictions.csv",
        "t3": hybrid_dir / "xgb_cpi_t3_predictions.csv",
        "t5": hybrid_dir / "xgb_cpi_t5_predictions.csv",
    }

    hybrid_preds_by_h: dict[str, pd.DataFrame] = {}
    for h, path in hybrid_files.items():
        dfh = pd.read_csv(path)
        dfh = dfh.rename(columns={"dates": "date", "predicted": f"hybrid_pred_{h}"})
        dfh["date"] = pd.to_datetime(dfh["date"], errors="coerce")
        dfh = dfh.dropna(subset=["date"]).sort_values("date")[["date", f"hybrid_pred_{h}"]]
        hybrid_preds_by_h[h] = dfh

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "rf_vs_hybrid_actual_vs_predicted_test.png"

    horizon_specs = [
        ("t+1", "t1", "cpi_t1", "rf_pred_t1"),
        ("t+3", "t3", "cpi_t3", "rf_pred_t3"),
        ("t+5", "t5", "cpi_t5", "rf_pred_t5"),
    ]

    ACTUAL_STYLE = {"color": "black", "linewidth": 2.8, "alpha": 1.0}
    RF_STYLE = {"color": "tab:blue", "linewidth": 2.2, "alpha": 0.9, "linestyle": "--"}
    HY_STYLE = {"color": "tab:orange", "linewidth": 2.6, "alpha": 0.95, "linestyle": "-"}

    fig, axes = plt.subplots(3, 1, figsize=(12, 9), sharex=True)
    for ax, (label, h, actual_col, rf_pred_col) in zip(axes, horizon_specs):
        df_cmp = (
            rf_actuals[["date", actual_col]]
            .merge(rf_preds[["date", rf_pred_col]], on="date", how="inner")
            .merge(hybrid_preds_by_h[h], on="date", how="inner")
            .dropna()
            .reset_index(drop=True)
        )
        dates = pd.to_datetime(df_cmp["date"])
        ax.plot(dates, df_cmp[actual_col], label="Actual", **ACTUAL_STYLE)
        ax.plot(dates, df_cmp[rf_pred_col], label="RF", **RF_STYLE)
        ax.plot(dates, df_cmp[f"hybrid_pred_{h}"], label="Hybrid", **HY_STYLE)
        ax.set_title(label)
        ax.grid(True, alpha=0.25)
        ax.legend(loc="upper right")

    for ax in axes:
        ax.xaxis.set_major_locator(mdates.YearLocator())
        ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y"))
    plt.setp(axes[-1].xaxis.get_majorticklabels(), rotation=45, ha="right")
    plt.suptitle("RF vs Hybrid — Actual vs Predicted (Test)")
    plt.tight_layout()
    plt.savefig(out_path, dpi=220)
    plt.close(fig)

    return out_path


def main() -> int:
    from docx import Document

    root = _project_root()

    # Key inputs
    eda_dir = root / "reports" / "eda_graphs"
    rf_plots_dir = root / "artifacts" / "rf_baseline_plots"
    eval_dir = root / "artifacts" / "evaluation"
    results_baselines = root / "results" / "baselines"

    # Outputs
    out_doc = root / "reports" / "Chapter4_Results_and_Conclusions_UWE.docx"
    generated_fig_dir = eval_dir / "generated"

    doc = Document()
    _set_uwe_style(doc)

    _fmt_heading(doc, "Chapter 4 — Results and Conclusions", level=1)
    doc.add_paragraph(
        "This chapter summarises the key exploratory findings (EDA) and presents a horizon-wise evaluation of baseline and hybrid forecasting models for UK CPI inflation. "
        "Results are reported for horizons t+1, t+3, and t+5 using MAE, RMSE, and R², alongside visual diagnostics." 
    )

    # -------------------
    # 4.1 EDA highlights
    # -------------------
    _fmt_heading(doc, "4.1 Exploratory Data Analysis (EDA) — key findings", level=2)
    doc.add_paragraph(
        "EDA figures are included to motivate feature selection and highlight the behaviour of CPI and candidate drivers (bank rate, exchange rates, petrol/oil price changes, unemployment, and GVA growth). "
        "The plots also illustrate correlation structure, lead–lag patterns, rolling relationships, and potential structural changes." 
    )

    eda_figs = _sorted_figs(list(eda_dir.glob("*.png")))
    for i, p in enumerate(eda_figs, start=1):
        _add_figure(doc, p, caption=f"Figure 4.1.{i}: {p.name}")

    # -------------------
    # 4.2 ARIMA baseline
    # -------------------
    _fmt_heading(doc, "4.2 Statistical baseline — ARIMA", level=2)
    doc.add_paragraph(
        "ARIMA is treated as a classical univariate benchmark. Forecasts are loaded from saved artifacts and evaluated on the test period." 
    )

    # ARIMA metrics (single series)
    dates = np.load(results_baselines / "arima_test_dates.npy", allow_pickle=True)
    actual = np.load(results_baselines / "arima_test_actual.npy", allow_pickle=True)
    forecast = np.load(results_baselines / "arima_test_forecast.npy", allow_pickle=True)

    arima_mae = float(np.mean(np.abs(actual - forecast)))
    arima_rmse = float(np.sqrt(np.mean((actual - forecast) ** 2)))
    # R2
    ss_res = float(np.sum((actual - forecast) ** 2))
    ss_tot = float(np.sum((actual - np.mean(actual)) ** 2))
    arima_r2 = float(1.0 - ss_res / ss_tot) if ss_tot > 0 else float("nan")

    arima_df = pd.DataFrame(
        [{"split": "test", "series": "CPI (univariate)", "n": int(len(actual)), "MAE": arima_mae, "RMSE": arima_rmse, "R2": arima_r2}]
    )
    _add_df_table(doc, arima_df.round(3), title="ARIMA test-set metrics")

    arima_plot = _make_arima_plot(root, generated_fig_dir)
    _add_figure(doc, arima_plot, caption="Figure 4.2.1: ARIMA baseline — Actual vs Forecast (test)")

    # --------------------------------
    # 4.3 Forecast performance summary
    # --------------------------------
    _fmt_heading(doc, "4.3 Forecast performance summary (test set)", level=2)
    doc.add_paragraph(
        "Table 4.3.1 compares the principal models (ARIMA, Random Forest, Hybrid LSTM–XGBoost) on the common test-set dates for each horizon." 
    )

    t1 = _safe_read_csv(eval_dir / "table1_forecast_performance_comparison.csv")
    t1_fmt = t1.copy()
    t1_fmt["horizon"] = t1_fmt["horizon"].map({1: "t+1", 3: "t+3", 5: "t+5"}).fillna(t1_fmt["horizon"].astype(str))
    _add_df_table(doc, t1_fmt.round(3), title="Table 4.3.1: Forecast performance comparison (test set)")

    # Add evaluation figures (already produced)
    eval_pngs = [
        *sorted(eval_dir.glob("graph1_actual_vs_predicted_h*.png")),
        *sorted(eval_dir.glob("forecast_diagnostics_h*.png")),
        *sorted(eval_dir.glob("error_bars_h*.png")),
        *sorted(eval_dir.glob("graph5_error_distribution_h*.png")),
        *sorted(eval_dir.glob("graph6_rolling_rmse_h*.png")),
    ]

    for i, p in enumerate(eval_pngs, start=1):
        _add_figure(doc, p, caption=f"Figure 4.3.{i}: {p.name}")

    # -------------------
    # 4.4 Random Forest
    # -------------------
    _fmt_heading(doc, "4.4 Machine learning baseline — Random Forest (RF)", level=2)
    doc.add_paragraph(
        "The RF baseline is evaluated across train/validation/test and provides multi-horizon predictions. Diagnostic plots are provided for each split, plus combined 3-in-1 horizon images." 
    )

    rf_metrics = _safe_read_csv(root / "artifacts" / "rf_baseline_metrics.csv")
    _add_df_table(doc, rf_metrics.round(3), title="Table 4.4.1: RF baseline metrics (train/validation/test)")

    # Prefer the combined horizon plots
    rf_figs = [
        rf_plots_dir / "scatter_train_all_horizons.png",
        rf_plots_dir / "line_train_all_horizons.png",
        rf_plots_dir / "scatter_validation_all_horizons.png",
        rf_plots_dir / "line_validation_all_horizons.png",
        rf_plots_dir / "scatter_test_all_horizons.png",
        rf_plots_dir / "line_test_all_horizons.png",
    ]

    for i, p in enumerate(rf_figs, start=1):
        _add_figure(doc, p, caption=f"Figure 4.4.{i}: {p.name}")

    # -------------------
    # 4.5 Hybrid model
    # -------------------
    _fmt_heading(doc, "4.5 Proposed model — Hybrid LSTM–XGBoost", level=2)
    doc.add_paragraph(
        "The Hybrid model combines an LSTM representation stage with an XGBoost meta-learner. Test-set results are reported by horizon and compared directly to RF." 
    )

    # Include project report figures if present
    report_figs = [
        root / "reports" / "stacking_ensemble_comparison.png",
        root / "reports" / "xgboost_actual_vs_predicted.png",
        root / "reports" / "xgboost_feature_importance.png",
        root / "reports" / "feature_importance_meta_learner.png",
        root / "reports" / "xgboost_results.png",
    ]

    for i, p in enumerate([p for p in report_figs if p.exists()], start=1):
        _add_figure(doc, p, caption=f"Figure 4.5.{i}: {p.name}")

    # RF vs Hybrid comparison plot (generated)
    rfhy_plot = _make_rf_vs_hybrid_plot(root, generated_fig_dir)
    _add_figure(doc, rfhy_plot, caption="Figure 4.5.99: RF vs Hybrid — Actual vs Predicted (test)")

    # -------------------
    # 4.6 Conclusions
    # -------------------
    _fmt_heading(doc, "4.6 Conclusions", level=2)
    doc.add_paragraph(
        "Key conclusions from the empirical results are summarised below." 
    )

    conclusions = [
        "Forecast error increases with horizon length (t+1 → t+5), which is consistent with compounding uncertainty in macroeconomic forecasting.",
        "Across the common test window, Random Forest provides the strongest overall baseline among the evaluated methods in this project’s artifact set.",
        "The Hybrid LSTM–XGBoost model demonstrates competitive tracking in parts of the test period, but its aggregate error metrics are not consistently better than RF on the common-date comparison.",
        "ARIMA remains a useful classical benchmark but is generally outperformed by feature-based ML baselines on the selected test window.",
        "Future work should explore: re-tuning hybrid components, alternative sequence models, probabilistic forecasting, and robustness checks over multiple test windows / rolling-origin evaluation.",
    ]

    for c in conclusions:
        doc.add_paragraph(c, style=None).paragraph_format.left_indent = None

    # Apply line spacing to paragraphs created before style enforcement
    _set_paragraph_spacing(doc)

    out_doc.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_doc)
    print(f"Saved: {out_doc}")
    print(f"Generated figures in: {generated_fig_dir}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
